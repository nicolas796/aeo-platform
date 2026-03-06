from flask import Blueprint, render_template, jsonify, request, redirect, url_for, flash, send_file, current_app
from flask_login import login_required, current_user
from app.models import db, WeeklyReport, ContentSuggestion, GeneratedContent, ContentShare
from app.services.content_generation import ContentGenerationService
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re
import io
import secrets

reports_bp = Blueprint('reports', __name__)

@reports_bp.route('/')
@login_required
def index():
    tenant_id = current_user.tenant_id
    reports = WeeklyReport.query.filter_by(tenant_id=tenant_id).order_by(WeeklyReport.report_date.desc()).all()
    return render_template('reports/index.html', reports=reports)

@reports_bp.route('/<int:id>')
@login_required
def detail(id):
    tenant_id = current_user.tenant_id
    report = WeeklyReport.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    return render_template('reports/detail.html', report=report)

@reports_bp.route('/suggestions')
@login_required
def suggestions():
    tenant_id = current_user.tenant_id
    suggestions = ContentSuggestion.query.filter_by(tenant_id=tenant_id).order_by(ContentSuggestion.created_at.desc()).all()
    return render_template('reports/suggestions.html', suggestions=suggestions)

import threading

@reports_bp.route('/suggestions/<int:id>/approve', methods=['POST'])
@login_required
def approve_suggestion(id):
    """Approve a content suggestion and start generation in background"""
    tenant_id = current_user.tenant_id
    suggestion = ContentSuggestion.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    
    # Check if already generated
    if suggestion.status == 'created':
        generated = GeneratedContent.query.filter_by(suggestion_id=suggestion.id).first()
        if generated:
            return redirect(url_for('reports.edit_content', id=generated.id))
    
    # Check if already generating
    if suggestion.status == 'generating':
        return redirect(url_for('reports.generating', id=id))
    
    # Mark as generating and start background thread
    suggestion.status = 'generating'
    db.session.commit()
    
    # Capture app object before starting thread
    from flask import current_app
    app = current_app._get_current_object()
    suggestion_id = suggestion.id
    
    def generate_async():
        with app.app_context():
            try:
                service = ContentGenerationService()
                generated = service.generate_content(suggestion_id)
                print(f"Content generated: {generated.title}")
            except Exception as e:
                print(f"Generation error: {e}")
                # Re-fetch suggestion in this context
                from app.models import ContentSuggestion
                s = ContentSuggestion.query.get(suggestion_id)
                if s:
                    s.status = 'pending'
                    from app.models import db
                    db.session.commit()
    
    thread = threading.Thread(target=generate_async)
    thread.daemon = True
    thread.start()
    
    return redirect(url_for('reports.generating', id=id))

@reports_bp.route('/suggestions/<int:id>/generating')
@login_required
def generating(id):
    """Show generating page while content is being created"""
    tenant_id = current_user.tenant_id
    suggestion = ContentSuggestion.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    
    # Check if done
    if suggestion.status == 'created':
        generated = GeneratedContent.query.filter_by(suggestion_id=suggestion.id).first()
        if generated:
            return redirect(url_for('reports.edit_content', id=generated.id))
    
    return render_template('reports/generating.html', suggestion=suggestion)

@reports_bp.route('/suggestions/<int:id>/reject', methods=['POST'])
@login_required
def reject_suggestion(id):
    """Reject a content suggestion"""
    tenant_id = current_user.tenant_id
    suggestion = ContentSuggestion.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    
    suggestion.status = 'rejected'
    db.session.commit()
    
    flash('Suggestion rejected.', 'info')
    return redirect(url_for('reports.suggestions'))

@reports_bp.route('/content')
@login_required
def content_library():
    """Show all generated content"""
    tenant_id = current_user.tenant_id
    contents = GeneratedContent.query.filter_by(tenant_id=tenant_id).order_by(GeneratedContent.created_at.desc()).all()
    return render_template('reports/content_library.html', contents=contents)

@reports_bp.route('/content/<int:id>')
@login_required
def view_content(id):
    """View generated content"""
    tenant_id = current_user.tenant_id
    content = GeneratedContent.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    return render_template('reports/view_content.html', content=content)

@reports_bp.route('/content/<int:id>/edit')
@login_required
def edit_content(id):
    """Edit generated content"""
    tenant_id = current_user.tenant_id
    content = GeneratedContent.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    return render_template('reports/edit_content.html', content=content)

@reports_bp.route('/content/<int:id>/update', methods=['POST'])
@login_required
def update_content(id):
    """Save edited content"""
    tenant_id = current_user.tenant_id
    content = GeneratedContent.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    
    content.title = request.form.get('title', content.title)
    content.content = request.form.get('content', content.content)
    content.meta_description = request.form.get('meta_description', content.meta_description)
    content.word_count = len(content.content.split())
    content.status = request.form.get('status', content.status)
    
    db.session.commit()
    
    flash('Content saved successfully!', 'success')
    return redirect(url_for('reports.edit_content', id=id))

@reports_bp.route('/content/<int:id>/send', methods=['POST'])
@login_required
def send_content_email(id):
    """Send content teaser email with share link"""
    tenant_id = current_user.tenant_id
    content = GeneratedContent.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()

    to_email = request.form.get('to_email')
    message = request.form.get('message', '')

    if not to_email:
        flash('Please provide an email address.', 'error')
        return redirect(url_for('reports.edit_content', id=id))

    # Create a share link (expires in 30 days)
    token = secrets.token_urlsafe(32)
    share = ContentShare(
        content_id=content.id,
        shared_by=current_user.id,
        token=token,
        recipient_email=to_email,
        expires_at=datetime.utcnow() + timedelta(days=30)
    )
    db.session.add(share)
    db.session.commit()

    app_url = current_app.config.get('APP_URL')
    share_path = url_for('reports.view_shared_content', token=token)
    share_url = f"{app_url}{share_path}" if app_url else url_for('reports.view_shared_content', token=token, _external=True)

    from app.services.email_service import EmailService
    service = EmailService()
    success, error = service.send_content_for_review(
        to_email=to_email,
        content=content,
        share_url=share_url,
        sender_name=f"{current_user.first_name} {current_user.last_name}",
        message=message
    )
    if success:
        flash(f'Content shared with {to_email}!', 'success')
    else:
        flash(f'Failed to send email: {error}', 'error')

    return redirect(url_for('reports.edit_content', id=id))

@reports_bp.route('/content/<int:id>/export/word')
@login_required
def export_word(id):
    """Export content as formatted Word document"""
    tenant_id = current_user.tenant_id
    content = GeneratedContent.query.filter_by(id=id, tenant_id=tenant_id).first_or_404()
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(content.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add meta description as subtitle
    if content.meta_description:
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Meta Description: {content.meta_description}")
        meta_run.italic = True
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()  # Spacing
    
    # Convert markdown to HTML then parse for Word
    md_content = content.content
    
    # Simple markdown parsing for Word
    lines = md_content.split('\n')
    current_list = None
    
    for line in lines:
        line = line.strip()
        if not line:
            current_list = None
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            if current_list != 'bullet':
                current_list = 'bullet'
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            if current_list != 'number':
                current_list = 'number'
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
        # Regular paragraph with bold/italic support
        else:
            current_list = None
            p = doc.add_paragraph()
            
            # Parse inline formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('__') and part.endswith('__'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('_') and part.endswith('_'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
    
    # Add SEO section at the end
    doc.add_page_break()
    doc.add_heading('SEO Information', level=2)
    
    if content.seo_keyphrase:
        p = doc.add_paragraph()
        p.add_run('Primary Keyphrase: ').bold = True
        p.add_run(content.seo_keyphrase)
    
    # Internal links
    internal_links = content.get_internal_links()
    if internal_links:
        doc.add_heading('Suggested Internal Links', level=3)
        for link in internal_links:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'"{link["anchor_text"]}"').italic = True
            p.add_run(f' → {link["target_url"]}')
            if link.get('reason'):
                p = doc.add_paragraph(f'Reason: {link["reason"]}', style='List Bullet 2')
    
    # External links
    external_links = content.get_external_links()
    if external_links:
        doc.add_heading('Suggested External Links', level=3)
        for link in external_links:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(link.get('source_name', 'Source')).bold = True
            p.add_run(f': "{link["anchor_text"]}" → {link["target_url"]}')
            if link.get('reason'):
                p = doc.add_paragraph(f'Reason: {link["reason"]}', style='List Bullet 2')
    
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Generate filename
    filename = re.sub(r'[^\w\s-]', '', content.title).strip().replace(' ', '_') + '.docx'
    
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@reports_bp.route('/shared/<token>')
def view_shared_content(token):
    """Public page to view shared content (no login required)"""
    share = ContentShare.query.filter_by(token=token).first()

    if not share or share.is_expired():
        flash('This share link is invalid or has expired.', 'error')
        return redirect(url_for('auth.login'))

    content = share.content
    import bleach
    raw_html = markdown.markdown(content.content, extensions=['tables', 'fenced_code'])
    allowed_tags = list(bleach.ALLOWED_TAGS) + [
        'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'br', 'hr',
        'pre', 'code', 'table', 'thead', 'tbody', 'tr', 'th', 'td',
        'ul', 'ol', 'li', 'img', 'div', 'span', 'strong', 'em',
    ]
    allowed_attrs = dict(bleach.ALLOWED_ATTRIBUTES)
    allowed_attrs['img'] = ['src', 'alt', 'title']
    allowed_attrs['td'] = ['align']
    allowed_attrs['th'] = ['align']
    html_body = bleach.clean(raw_html, tags=allowed_tags, attributes=allowed_attrs)

    return render_template('reports/shared_content.html',
                           content=content, html_body=html_body, share=share)


@reports_bp.route('/shared/<token>/download')
def export_shared_word(token):
    """Download Word document for shared content (no login required)"""
    share = ContentShare.query.filter_by(token=token).first()

    if not share or share.is_expired():
        flash('This share link is invalid or has expired.', 'error')
        return redirect(url_for('auth.login'))

    content = share.content

    doc = Document()
    title = doc.add_heading(content.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if content.meta_description:
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Meta Description: {content.meta_description}")
        meta_run.italic = True
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

    for line in content.content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = re.sub(r'[^\w\s-]', '', content.title).strip().replace(' ', '_') + '.docx'

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )