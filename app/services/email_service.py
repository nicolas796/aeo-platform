"""Email service for sending content via SendGrid"""
import os
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition
import base64
from flask import current_app


class EmailService:
    """Send content via SendGrid email"""

    def __init__(self):
        api_key = current_app.config.get('SENDGRID_API_KEY') or os.environ.get('SENDGRID_API_KEY')
        self.api_key = api_key.strip() if api_key else None
        self.from_email = current_app.config.get('SENDGRID_FROM_EMAIL', 'noreply@aeoplatform.local')
        self.client = None
        if self.api_key:
            self.client = SendGridAPIClient(self.api_key)
            print(f"EmailService initialized with API key: {self.api_key[:10]}...")
            print(f"From email: {self.from_email}")

    def send_content_for_review(self, to_email, content, message=None, include_word=True, include_thumbnail=True):
        """Send generated content to someone for review

        Args:
            to_email: Recipient email address
            content: GeneratedContent object
            message: Optional personal message
            include_word: Whether to attach Word document
            include_thumbnail: Whether to attach thumbnail image

        Returns:
            tuple: (success: bool, error_message: str or None)
        """
        if not self.client:
            return False, "SendGrid API key not configured"

        try:
            # Build email subject
            subject = f"[Review] {content.title}"

            # Build email body
            body_lines = [
                f"<h2>{content.title}</h2>",
                "<p>Please review the attached content.</p>"
            ]

            if message:
                body_lines.append(f"<p><strong>Message:</strong> {message}</p>")

            body_lines.extend([
                f"<p><strong>Word Count:</strong> {content.word_count}</p>",
                f"<p><strong>SEO Keyphrase:</strong> {content.seo_keyphrase or 'N/A'}</p>",
                "<hr>",
                "<h3>Content Preview:</h3>",
                f"<pre style='white-space: pre-wrap; background: #f5f5f5; padding: 15px; border-radius: 5px;'>{content.content[:2000]}...</pre>" if len(content.content) > 2000 else f"<pre style='white-space: pre-wrap; background: #f5f5f5; padding: 15px; border-radius: 5px;'>{content.content}</pre>",
                "<hr>",
                "<p><em>Sent from AEO Platform</em></p>"
            ])

            html_content = "\n".join(body_lines)

            # Build plain text version for deliverability
            plain_lines = [
                content.title,
                "",
                "Please review the attached content.",
            ]
            if message:
                plain_lines.extend(["", f"Message: {message}"])
            preview = content.content[:2000]
            plain_lines.extend([
                "",
                f"Word Count: {content.word_count}",
                f"SEO Keyphrase: {content.seo_keyphrase or 'N/A'}",
                "",
                "--- Content Preview ---",
                preview,
                "",
                "Sent from AEO Platform",
            ])
            plain_text = "\n".join(plain_lines)

            # Create email
            mail = Mail(
                from_email=self.from_email,
                to_emails=to_email,
                subject=subject,
                plain_text_content=plain_text,
                html_content=html_content
            )

            # Attach Word document if requested
            if include_word:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io
                import re

                doc = Document()

                # Add title
                title = doc.add_heading(content.title, level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add content
                lines = content.content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.\s', line):
                        text = re.sub(r'^\d+\.\s', '', line)
                        doc.add_paragraph(text, style='List Number')
                    else:
                        p = doc.add_paragraph()
                        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', line)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            elif part.startswith('*') and part.endswith('*'):
                                run = p.add_run(part[1:-1])
                                run.italic = True
                            elif part.startswith('__') and part.endswith('__'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            elif part.startswith('_') and part.endswith('_'):
                                run = p.add_run(part[1:-1])
                                run.italic = True
                            else:
                                p.add_run(part)

                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Attach
                filename = re.sub(r'[^\w\s-]', '', content.title).strip().replace(' ', '_') + '.docx'
                encoded = base64.b64encode(buffer.read()).decode()

                attachment = Attachment()
                attachment.file_content = FileContent(encoded)
                attachment.file_name = FileName(filename)
                attachment.file_type = FileType('application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                attachment.disposition = Disposition('attachment')
                mail.add_attachment(attachment)

            # Attach thumbnail if requested and exists
            if include_thumbnail and content.thumbnail_path:
                import os
                thumbnail_full_path = os.path.join(current_app.root_path, 'static', content.thumbnail_path)
                if os.path.exists(thumbnail_full_path):
                    with open(thumbnail_full_path, 'rb') as f:
                        thumbnail_data = f.read()

                    encoded_thumb = base64.b64encode(thumbnail_data).decode()
                    thumb_attachment = Attachment()
                    thumb_attachment.file_content = FileContent(encoded_thumb)
                    thumb_attachment.file_name = FileName('thumbnail.png')
                    thumb_attachment.file_type = FileType('image/png')
                    thumb_attachment.disposition = Disposition('attachment')
                    mail.add_attachment(thumb_attachment)

            # Send email
            response = self.client.send(mail)

            print(f"SendGrid response status: {response.status_code}")
            print(f"SendGrid response body: {response.body}")

            if response.status_code in [200, 201, 202]:
                return True, None
            else:
                return False, f"SendGrid returned status {response.status_code}: {response.body}"

        except Exception as e:
            import traceback
            print(f"SendGrid error: {e}")
            print(traceback.format_exc())
            error_msg = self._extract_error(e)
            return False, error_msg

    def send_invitation_email(self, to_email, inviter, invite_url):
        """Send team invitation email

        Args:
            to_email: Recipient email address
            inviter: User object who sent the invitation
            invite_url: Full URL to accept the invitation

        Returns:
            tuple: (success: bool, error_message: str or None)
        """
        if not self.client:
            return False, "SendGrid API key not configured"

        try:
            subject = f"You've been invited to join {inviter.tenant.name} on AEO Platform"

            html_content = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #4F46E5;">You're Invited!</h2>
                <p><strong>{inviter.first_name} {inviter.last_name}</strong> has invited you to join
                <strong>{inviter.tenant.name}</strong> on AEO Platform.</p>

                <p>AEO Platform helps brands optimize their visibility in AI-powered search engines.</p>

                <div style="margin: 30px 0; text-align: center;">
                    <a href="{invite_url}" style="display: inline-block; padding: 14px 32px; background-color: #4F46E5; color: #ffffff; text-decoration: none; border-radius: 8px; font-weight: bold;">Accept Invitation</a>
                </div>

                <p style="color: #666; font-size: 13px;">Or copy and paste this link into your browser:<br>
                <a href="{invite_url}" style="color: #4F46E5;">{invite_url}</a></p>

                <p style="color: #666; font-size: 14px; margin-top: 30px;">
                    This invitation will expire in 7 days.<br>
                    If you didn't expect this invitation, you can ignore this email.
                </p>
            </div>
            """

            # Create plain text version
            plain_text = f"""You've been invited to join {inviter.tenant.name} on AEO Platform.

{inviter.first_name} {inviter.last_name} has invited you to join their team.

To accept this invitation, visit the following link:
{invite_url}

This invitation will expire in 7 days.

If you didn't expect this invitation, you can ignore this email.
"""

            from sendgrid.helpers.mail import Content

            mail = Mail(
                from_email=self.from_email,
                to_emails=to_email,
                subject=subject,
                plain_text_content=plain_text,
                html_content=html_content
            )

            response = self.client.send(mail)

            if response.status_code in [200, 201, 202]:
                return True, None
            else:
                return False, f"SendGrid returned status {response.status_code}"

        except Exception as e:
            import traceback
            print(f"SendGrid invitation error: {e}")
            print(traceback.format_exc())
            error_msg = self._extract_error(e)
            return False, error_msg

    @staticmethod
    def _extract_error(e):
        """Extract detailed error message from SendGrid API exceptions."""
        status = getattr(e, 'status_code', None)
        body = getattr(e, 'body', None)
        if status and body:
            try:
                body_str = body.decode('utf-8') if isinstance(body, bytes) else str(body)
            except Exception:
                body_str = str(body)
            return f"HTTP {status}: {body_str}"
        return str(e)
